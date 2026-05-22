from __future__ import annotations

import hashlib
import re
from copy import deepcopy
from datetime import date, datetime
from pathlib import Path
from tempfile import NamedTemporaryFile

from django.core.exceptions import ValidationError
from django.utils import timezone

from .docx_template_models import DOCXTemplate, validate_docx_template_file
from .models import AuditLog, FormEntry, PDFDocument
from .pdf_services import get_private_document_root

PLACEHOLDER_RE = re.compile(r"{{\s*([a-zA-Z0-9_.:-]+)\s*}}")


def _require_docx():
    try:
        from docx import Document
    except ImportError as exc:  # pragma: no cover - dependency guidance only
        raise ValidationError(
            "python-docx ist nicht installiert. Bitte `poetry add python-docx` ausfuehren."
        ) from exc
    return Document


def _iter_paragraphs(document):
    for paragraph in document.paragraphs:
        yield paragraph
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph


def _paragraph_texts(document) -> list[str]:
    texts = []
    for paragraph in _iter_paragraphs(document):
        if paragraph.text:
            texts.append(paragraph.text)
    return texts


def extract_docx_placeholders(uploaded_file) -> dict:
    validate_docx_template_file(uploaded_file)
    Document = _require_docx()
    position = None
    if hasattr(uploaded_file, "tell") and hasattr(uploaded_file, "seek"):
        position = uploaded_file.tell()
        uploaded_file.seek(0)
    document = Document(uploaded_file)
    if position is not None:
        uploaded_file.seek(position)
    placeholders: list[str] = []
    for text in _paragraph_texts(document):
        for match in PLACEHOLDER_RE.finditer(text):
            key = match.group(1).strip()
            if key and key not in placeholders:
                placeholders.append(key)
    return {
        "placeholders": placeholders,
        "paragraph_count": len(document.paragraphs),
        "table_count": len(document.tables),
    }


def analyze_and_store_docx_template(template: DOCXTemplate) -> DOCXTemplate:
    with template.template_file.open("rb") as fh:
        analysis = extract_docx_placeholders(fh)
    template.placeholder_keys = analysis["placeholders"]
    template.analysis = analysis
    template.save(update_fields=["placeholder_keys", "analysis", "updated_at"])
    return template


def create_docx_template(
    *, form, uploaded_file, title: str, description: str = "", user=None
) -> DOCXTemplate:
    validate_docx_template_file(uploaded_file)
    analysis = extract_docx_placeholders(uploaded_file)
    template = DOCXTemplate.objects.create(
        form=form,
        title=title or getattr(uploaded_file, "name", "DOCX-Vorlage"),
        description=description,
        template_file=uploaded_file,
        original_filename=getattr(uploaded_file, "name", "template.docx")[:255],
        content_type=(
            getattr(uploaded_file, "content_type", "")
            or DOCXTemplate._meta.get_field("content_type").default
        ),
        file_size=int(getattr(uploaded_file, "size", 0) or 0),
        placeholder_keys=analysis["placeholders"],
        analysis=analysis,
        uploaded_by=user,
        created_by=user,
        updated_by=user,
    )
    AuditLog.objects.create(
        actor=user,
        event_type=AuditLog.EventType.CREATED,
        target_model="DOCXTemplate",
        target_id=template.pk,
        form=form,
        message="DOCX-Vorlage wurde hochgeladen.",
        metadata={"template_id": str(template.pk), "placeholders": template.placeholder_keys},
    )
    return template


def _format_value(value) -> str:
    if isinstance(value, datetime):
        return timezone.localtime(value).strftime("%d.%m.%Y %H:%M")
    if isinstance(value, date):
        return value.strftime("%d.%m.%Y")
    if isinstance(value, bool):
        return "Ja" if value else "Nein"
    if isinstance(value, list):
        if value and all(isinstance(item, dict) for item in value):
            rows = []
            for item in value:
                rows.append(", ".join(f"{key}: {_format_value(val)}" for key, val in item.items()))
            return "\n".join(rows)
        return ", ".join(_format_value(item) for item in value)
    if isinstance(value, dict):
        if value.get("filename"):
            return str(value.get("filename"))
        return ", ".join(f"{key}: {_format_value(val)}" for key, val in value.items())
    if value in (None, ""):
        return ""
    return str(value)


def build_docx_context(form_entry: FormEntry) -> dict[str, str]:
    data = deepcopy(form_entry.data or {})
    bewohner = form_entry.bewohner
    context = {
        "entry_id": str(form_entry.public_id),
        "form_title": form_entry.form.title,
        "form_key": form_entry.form.key,
        "status": form_entry.get_status_display(),
        "datum": timezone.localtime(timezone.now()).strftime("%d.%m.%Y"),
        "bewohner_name": str(bewohner),
        "bewohner_vorname": getattr(bewohner, "first_name", "") or "",
        "bewohner_nachname": getattr(bewohner, "last_name", "") or "",
        "bewohner_nummer": getattr(bewohner, "resident_number", "") or "",
        "bewohner_zimmer": getattr(bewohner, "room_label", "") or "",
    }
    if getattr(bewohner, "date_of_birth", None):
        context["bewohner_geburtsdatum"] = _format_value(bewohner.date_of_birth)
    for key, value in data.items():
        context[str(key)] = _format_value(value)
    return context


def _replace_in_paragraph(paragraph, context: dict[str, str]) -> None:
    if "{{" not in paragraph.text:
        return
    original = paragraph.text
    replaced = PLACEHOLDER_RE.sub(lambda match: context.get(match.group(1).strip(), ""), original)
    if replaced == original:
        return
    if paragraph.runs:
        paragraph.runs[0].text = replaced
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(replaced)


def fill_docx_template_bytes(*, template: DOCXTemplate, form_entry: FormEntry) -> bytes:
    Document = _require_docx()
    context = build_docx_context(form_entry)
    with template.template_file.open("rb") as fh:
        document = Document(fh)
    for paragraph in _iter_paragraphs(document):
        _replace_in_paragraph(paragraph, context)
    with NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp_path = Path(tmp.name)
    try:
        document.save(tmp_path)
        return tmp_path.read_bytes()
    finally:
        try:
            tmp_path.unlink(missing_ok=True)
        except Exception:
            pass


def get_default_docx_template(form) -> DOCXTemplate | None:
    return (
        DOCXTemplate.objects.filter(
            form=form,
            is_default=True,
            status=DOCXTemplate.TemplateStatus.ACTIVE,
        )
        .order_by("-activated_at", "-created_at")
        .first()
    )


def generate_docx_document(
    *, form_entry: FormEntry, template: DOCXTemplate | None = None, user=None
) -> PDFDocument:
    template = template or get_default_docx_template(form_entry.form)
    if not template:
        raise ValidationError("Fuer dieses Formular ist keine aktive DOCX-Vorlage hinterlegt.")
    docx_bytes = fill_docx_template_bytes(template=template, form_entry=form_entry)
    sha256 = hashlib.sha256(docx_bytes).hexdigest()
    now = timezone.now()
    storage_key = (
        f"docx_documents/{form_entry.pk}/{now.strftime('%Y%m%d_%H%M%S')}_{sha256[:12]}.docx"
    )
    target_path = (get_private_document_root() / storage_key).resolve()
    target_path.parent.mkdir(parents=True, exist_ok=True)
    target_path.write_bytes(docx_bytes)
    document = PDFDocument.objects.create(
        form=form_entry.form,
        form_entry=form_entry,
        bewohner=form_entry.bewohner,
        document_kind=PDFDocument.DocumentKind.FINAL,
        status=PDFDocument.GenerationStatus.GENERATED,
        storage_key=storage_key,
        original_filename=f"{form_entry.form.key}_{form_entry.public_id}.docx",
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        file_size=len(docx_bytes),
        sha256=sha256,
        generated_at=now,
        access_policy={
            "private": True,
            "generated_by": str(user.pk) if user else None,
            "download_requires_permission": True,
            "source_template_id": str(template.pk),
            "document_format": "docx",
        },
        created_by=user,
        updated_by=user,
    )
    AuditLog.objects.create(
        actor=user,
        event_type=AuditLog.EventType.PDF_RENDERED,
        target_model="PDFDocument",
        target_id=document.pk,
        bewohner=form_entry.bewohner,
        form=form_entry.form,
        form_entry=form_entry,
        message="DOCX-Dokument wurde aus Vorlage erzeugt und privat gespeichert.",
        metadata={
            "template_id": str(template.pk),
            "document_id": str(document.pk),
            "sha256": sha256,
        },
    )
    return document
