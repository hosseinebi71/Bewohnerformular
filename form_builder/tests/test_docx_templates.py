from __future__ import annotations

import tempfile
from io import BytesIO

from django.contrib.auth import get_user_model
from django.core.exceptions import ValidationError
from django.core.files.uploadedfile import SimpleUploadedFile
from django.test import TestCase, override_settings

from form_builder.docx_template_models import DOCXTemplate
from form_builder.docx_template_services import (
    build_docx_context,
    create_docx_template,
    fill_docx_template_bytes,
    generate_docx_document,
)
from form_builder.models import Bewohner, Field, Form, FormEntry


def _make_docx_upload(text: str = "Name: {{bewohner_name}} / Grund: {{grund}}"):
    from docx import Document

    document = Document()
    document.add_paragraph(text)
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Datum"
    table.cell(0, 1).text = "{{datum}}"
    payload = BytesIO()
    document.save(payload)
    payload.seek(0)
    return SimpleUploadedFile(
        "template.docx",
        payload.read(),
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


@override_settings(MEDIA_ROOT=tempfile.mkdtemp(), PRIVATE_DOCUMENT_ROOT=tempfile.mkdtemp())
class DOCXTemplateTests(TestCase):
    def setUp(self):
        User = get_user_model()
        self.user = User.objects.create_user(username="admin", password="test", is_staff=True)
        self.form = Form.objects.create(
            key="maengelmeldung",
            version=1,
            title="Maengelm eldung".replace(" ", ""),
            status=Form.PublicationStatus.DRAFT,
            created_by=self.user,
            updated_by=self.user,
        )
        self.field = Field.objects.create(
            form=self.form,
            key="grund",
            label="Grund",
            field_type=Field.FieldType.TEXT,
            position=1,
            created_by=self.user,
            updated_by=self.user,
        )
        self.form.sync_schema()
        self.bewohner = Bewohner.objects.create(
            resident_number="B-1",
            first_name="Max",
            last_name="Muster",
            room_label="101",
            created_by=self.user,
            updated_by=self.user,
        )
        self.entry = FormEntry.objects.create(
            form=self.form,
            bewohner=self.bewohner,
            form_snapshot=self.form.schema,
            data={"grund": "Defekter Wasserhahn"},
            created_by=self.user,
            updated_by=self.user,
        )

    def test_docx_upload_extracts_placeholders(self):
        template = create_docx_template(
            form=self.form,
            uploaded_file=_make_docx_upload(),
            title="Standard DOCX",
            user=self.user,
        )
        self.assertIn("bewohner_name", template.placeholder_keys)
        self.assertIn("grund", template.placeholder_keys)
        self.assertIn("datum", template.placeholder_keys)

    def test_docx_generation_replaces_placeholders(self):
        template = create_docx_template(
            form=self.form,
            uploaded_file=_make_docx_upload(),
            title="Standard DOCX",
            user=self.user,
        )
        generated = fill_docx_template_bytes(template=template, form_entry=self.entry)
        self.assertGreater(len(generated), 1000)
        from docx import Document

        document = Document(BytesIO(generated))
        text = "\n".join(paragraph.text for paragraph in document.paragraphs)
        self.assertIn("Defekter Wasserhahn", text)
        self.assertNotIn("{{grund}}", text)

    def test_generated_docx_is_stored_as_private_output_document(self):
        template = create_docx_template(
            form=self.form,
            uploaded_file=_make_docx_upload(),
            title="Standard DOCX",
            user=self.user,
        )
        template.activate(user=self.user)
        document = generate_docx_document(form_entry=self.entry, user=self.user)
        self.assertEqual(
            document.content_type,
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        self.assertTrue(document.storage_key.endswith(".docx"))
        self.assertEqual(document.access_policy["source_template_id"], str(template.pk))

    def test_context_contains_resident_and_field_values(self):
        context = build_docx_context(self.entry)
        self.assertIn("Muster", context["bewohner_name"])
        self.assertEqual(context["grund"], "Defekter Wasserhahn")

    def test_rejects_macro_enabled_files(self):
        upload = SimpleUploadedFile(
            "template.docm",
            b"not really a docx",
            content_type="application/vnd.ms-word.document.macroEnabled.12",
        )
        with self.assertRaises(ValidationError):
            create_docx_template(form=self.form, uploaded_file=upload, title="Bad", user=self.user)

    def test_template_model_can_be_activated(self):
        template = create_docx_template(
            form=self.form,
            uploaded_file=_make_docx_upload(),
            title="Standard DOCX",
            user=self.user,
        )
        template.activate(user=self.user)
        template.refresh_from_db()
        self.assertEqual(template.status, DOCXTemplate.TemplateStatus.ACTIVE)
        self.assertTrue(template.is_default)
